from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path("/root/synapse")
TOOLS = ROOT / "tools"
WEB_ADMIN = ROOT / "backend/app/web_admin"
ASSETS_V2 = TOOLS / "admin-manual-assets-v2"


def build_text(now: str) -> str:
    return f"""# Инструкция администратора Алхимик v3

Актуально на: {now}

## 1) Вход и старт смены
- Откройте `/api/v1/admin/web`.
- Нажмите `Вход в систему` в правом верхнем углу.
- Войдите по логину/паролю или OTP.
- Перейдите во вкладку `Безопасность` и нажмите `Утренний чек`.

## 2) Типовые ошибки и SLA реакции
| Ситуация | Действие | SLA |
|---|---|---|
| High алерт безопасности | Обновить алерты, добавить комментарий, подтвердить/снять, эскалировать owner | до 15 минут |
| Ошибка теста восстановления | Перезапустить тест, проверить практические действия и SLA | до 30 минут |
| Ошибка mobile smoke | Сохранить результат, проверить mobile readiness и content ingestion | до 30 минут |
| UI не реагирует | Ctrl+F5, повторный вход, проверить runbook статусы | до 10 минут |

## 3) Runbook сценарии
### Утренний чек
1. Нажать `Утренний чек`.
2. Проверить `runbookActionStatus`, `restoreActionStatus`, `alertsActionStatus`, `mobileActionStatus`.
3. При отклонениях выгрузить JSON/CSV.

### Инцидент-проверка
1. Нажать `Инцидент-проверка`.
2. Проверить практические действия, алерты, backup history, mobile/content readiness.
3. Зафиксировать выгрузку и передать ответственным.

## 4) Чеклист передачи смены
- [ ] Обработаны все high-алерты.
- [ ] Проверен статус теста восстановления.
- [ ] Проверены mobile/content readiness.
- [ ] При отклонениях выгружен отчет (JSON/CSV).
- [ ] Переданы комментарии следующему администратору.
- [ ] Выполнен выход (`Выйти`).

## 5) Релизный Go/No-Go
- Во вкладке `Безопасность` нажмите `Go/No-Go`.
- Статус `GO` означает, что обязательные гейты (security/payments/mobile/content/docs/high-alerts) пройдены.
- Статус `NO-GO` означает, что есть блокеры, релиз откладывается до устранения.

## 6) Printable one-page v3
- Откройте `/api/v1/admin/web/assets/admin-one-page-ru-v3.html`.
- Поля `Дата смены`, `Время передачи`, `Сдающий`, `Принимающий` заполняются автоматически/вручную.
- Поле `Инциденты за смену` заполняется перед печатью.
"""


def build_docx(now: str, out_path: Path) -> None:
    shots = [
        ASSETS_V2 / "01-login-modal.png",
        ASSETS_V2 / "02-users-dashboard.png",
        ASSETS_V2 / "03-security-dashboard.png",
        ASSETS_V2 / "04-guide-overlay.png",
    ]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Инструкция администратора Алхимик v3", level=0)
    doc.add_paragraph(f"Актуально на: {now}")

    doc.add_heading("Скрин 1: Вход в систему", level=1)
    if shots[0].exists():
        doc.add_picture(str(shots[0]), width=Inches(7.0))

    doc.add_heading("Скрин 2: Рабочая панель", level=1)
    if shots[1].exists():
        doc.add_picture(str(shots[1]), width=Inches(7.0))

    doc.add_heading("Скрин 3: Вкладка Безопасность", level=1)
    if shots[2].exists():
        doc.add_picture(str(shots[2]), width=Inches(7.0))

    doc.add_heading("Скрин 4: Инструкция в веб-админке", level=1)
    if shots[3].exists():
        doc.add_picture(str(shots[3]), width=Inches(7.0))

    doc.add_heading("SLA и действия", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Ситуация"
    hdr[1].text = "Действие"
    hdr[2].text = "SLA"
    rows = [
        ("High алерт", "Подтвердить/снять, комментарий, эскалация", "до 15 минут"),
        ("Ошибка backup test", "Перезапуск + проверка действий", "до 30 минут"),
        ("Ошибка mobile smoke", "Сохранить, проверить readiness", "до 30 минут"),
        ("UI не отвечает", "Ctrl+F5, вход заново", "до 10 минут"),
    ]
    for a, b, c in rows:
        r = table.add_row().cells
        r[0].text = a
        r[1].text = b
        r[2].text = c

    doc.add_heading("Чеклист передачи смены", level=1)
    for line in [
        "Обработаны high-алерты",
        "Проверен статус теста восстановления",
        "Проверены mobile/content readiness",
        "Выгружен отчет при отклонениях",
        "Передан комментарий следующей смене",
        "Выполнен выход из системы",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Подписи передачи смены", level=1)
    doc.add_paragraph("Сдающий смену: ____________________")
    doc.add_paragraph("Принимающий смену: ____________________")
    doc.add_paragraph("Время передачи: ____________________")

    doc.save(str(out_path))


def build_one_page_html(now: str, out_path: Path) -> None:
    html = f"""<!doctype html>
<html lang="ru">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>Алхимик One-Page v3</title>
  <style>
    body {{ font-family: Arial, sans-serif; color: #1f2f35; margin: 20px; }}
    h1 {{ margin: 0 0 8px; }}
    .muted {{ color: #5e727a; margin-bottom: 12px; }}
    .meta {{ display: grid; grid-template-columns: 1fr 1fr; gap: 8px; margin-bottom: 12px; }}
    .meta3 {{ display: grid; grid-template-columns: 1fr 1fr 1fr; gap: 8px; margin-bottom: 12px; }}
    .meta input {{ width: 100%; padding: 6px; border: 1px solid #c5d5dc; border-radius: 6px; }}
    table {{ width: 100%; border-collapse: collapse; margin-bottom: 12px; }}
    th, td {{ border: 1px solid #c5d5dc; padding: 7px; text-align: left; font-size: 13px; }}
    .check {{ margin: 6px 0; font-size: 14px; }}
    textarea {{ width: 100%; min-height: 120px; border: 1px solid #c5d5dc; border-radius: 6px; padding: 8px; font-family: inherit; }}
    @media print {{
      body {{ margin: 10mm; }}
      .noprint {{ display: none; }}
      input, textarea {{ border: none; padding: 0; }}
    }}
  </style>
</head>
<body>
  <h1>Алхимик: One-Page бланк дежурного администратора (v3)</h1>
  <div class="muted">Актуально на: {now}</div>
  <div class="meta">
    <label>Дата смены <input id="shiftDate" type="date" /></label>
    <label>ФИО дежурного <input id="dutyName" type="text" placeholder="Иванов И.И." /></label>
  </div>
  <div class="meta3">
    <label>Время передачи <input id="handoverTime" type="time" /></label>
    <label>Сдающий смену <input id="handoverOutgoing" type="text" placeholder="Петров П.П." /></label>
    <label>Принимающий смену <input id="handoverIncoming" type="text" placeholder="Иванова И.И." /></label>
  </div>

  <table>
    <thead><tr><th>Ситуация</th><th>Действие</th><th>SLA</th></tr></thead>
    <tbody>
      <tr><td>High алерт</td><td>Подтвердить/снять + комментарий + эскалация</td><td>до 15 минут</td></tr>
      <tr><td>Ошибка backup test</td><td>Повторный запуск + проверка практических действий</td><td>до 30 минут</td></tr>
      <tr><td>Ошибка mobile smoke</td><td>Сохранить результат + проверить readiness</td><td>до 30 минут</td></tr>
      <tr><td>UI не реагирует</td><td>Ctrl+F5 + повторный вход</td><td>до 10 минут</td></tr>
    </tbody>
  </table>

  <h3>Решение по релизу</h3>
  <div class="check">[ ] GO (обязательные гейты пройдены)</div>
  <div class="check">[ ] NO-GO (есть блокеры)</div>

  <h3>Чеклист передачи смены</h3>
  <div class="check">[ ] Обработаны high-алерты</div>
  <div class="check">[ ] Проверен тест восстановления</div>
  <div class="check">[ ] Проверены mobile/content readiness</div>
  <div class="check">[ ] При отклонениях выгружен отчет JSON/CSV</div>
  <div class="check">[ ] Комментарий передан следующему администратору</div>
  <div class="check">[ ] Выполнен выход из системы</div>

  <h3>Инциденты за смену</h3>
  <textarea id="incidents" placeholder="Опишите инциденты, время, действия и результат"></textarea>

  <div style="margin-top:12px;" class="noprint">
    <button onclick="window.print()">Печать</button>
  </div>

  <script>
    const dateEl = document.getElementById('shiftDate');
    const nameEl = document.getElementById('dutyName');
    const handoverTimeEl = document.getElementById('handoverTime');
    const outgoingEl = document.getElementById('handoverOutgoing');
    const incomingEl = document.getElementById('handoverIncoming');
    const incidentsEl = document.getElementById('incidents');
    const today = new Date();
    const iso = today.toISOString().slice(0, 10);
    if (!dateEl.value) dateEl.value = iso;
    const savedName = localStorage.getItem('synapse_duty_name') || '';
    const savedIncidents = localStorage.getItem('synapse_duty_incidents') || '';
    const savedHandoverTime = localStorage.getItem('synapse_handover_time') || '';
    const savedOutgoing = localStorage.getItem('synapse_handover_outgoing') || '';
    const savedIncoming = localStorage.getItem('synapse_handover_incoming') || '';
    if (savedName) nameEl.value = savedName;
    if (savedIncidents) incidentsEl.value = savedIncidents;
    if (savedHandoverTime) handoverTimeEl.value = savedHandoverTime;
    if (savedOutgoing) outgoingEl.value = savedOutgoing;
    if (savedIncoming) incomingEl.value = savedIncoming;
    nameEl.addEventListener('input', () => localStorage.setItem('synapse_duty_name', nameEl.value));
    incidentsEl.addEventListener('input', () => localStorage.setItem('synapse_duty_incidents', incidentsEl.value));
    handoverTimeEl.addEventListener('input', () => localStorage.setItem('synapse_handover_time', handoverTimeEl.value));
    outgoingEl.addEventListener('input', () => localStorage.setItem('synapse_handover_outgoing', outgoingEl.value));
    incomingEl.addEventListener('input', () => localStorage.setItem('synapse_handover_incoming', incomingEl.value));
  </script>
</body>
</html>
"""
    out_path.write_text(html, encoding="utf-8")


def main() -> None:
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    text = build_text(now)

    md = TOOLS / "admin-manual-ru-v3.md"
    txt = TOOLS / "admin-manual-ru-v3.txt"
    docx = TOOLS / "admin-manual-ru-v3.docx"
    one_page_html = TOOLS / "admin-one-page-ru-v3.html"

    md.write_text(text, encoding="utf-8")
    txt.write_text(text, encoding="utf-8")
    build_docx(now, docx)
    build_one_page_html(now, one_page_html)

    for src in [md, txt, docx, one_page_html]:
        (WEB_ADMIN / src.name).write_bytes(src.read_bytes())

    print("ok", md)
    print("ok", txt)
    print("ok", docx)
    print("ok", one_page_html)


if __name__ == "__main__":
    main()
