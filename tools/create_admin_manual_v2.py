from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from playwright.sync_api import sync_playwright


ROOT = Path("/root/synapse")
TOOLS = ROOT / "tools"
WEB_ADMIN = ROOT / "backend/app/web_admin"
ASSETS = TOOLS / "admin-manual-assets-v2"
ASSETS.mkdir(parents=True, exist_ok=True)


def capture_screenshots() -> list[Path]:
    result = []
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page(viewport={"width": 1600, "height": 2300})
        page.goto("http://127.0.0.1:8000/api/v1/admin/web", wait_until="networkidle")

        login_shot = ASSETS / "01-login-modal.png"
        page.click("#btnOpenAuth")
        page.wait_for_timeout(300)
        page.screenshot(path=str(login_shot), full_page=True)
        result.append(login_shot)

        page.fill("#adminLogin", "admin")
        page.fill("#adminPassword", "admin123")
        page.click("#btnAdminLogin")
        page.wait_for_timeout(1500)
        page.wait_for_selector("body.authorized")

        users_shot = ASSETS / "02-users-dashboard.png"
        page.screenshot(path=str(users_shot), full_page=True)
        result.append(users_shot)

        page.click('.viewBtn[data-view="security"]')
        page.wait_for_timeout(1000)
        security_shot = ASSETS / "03-security-dashboard.png"
        page.screenshot(path=str(security_shot), full_page=True)
        result.append(security_shot)

        page.click("#btnOpenGuide")
        page.wait_for_timeout(400)
        guide_shot = ASSETS / "04-guide-overlay.png"
        page.screenshot(path=str(guide_shot), full_page=True)
        result.append(guide_shot)

        browser.close()
    return result


def build_text(now: str) -> str:
    return f"""# Инструкция администратора Алхимик v2

Актуально на: {now}

## 1) Вход и старт смены
- Откройте `/api/v1/admin/web`.
- Нажмите `Вход в систему` в правом верхнем углу.
- Войдите по логину/паролю или OTP.
- Рекомендуемый старт: вкладка `Безопасность` -> `Утренний чек`.

## 2) Типовые ошибки и SLA реакции
| Ситуация | Действие | SLA |
|---|---|---|
| High алерт безопасности | Обновить алерты, добавить комментарий, подтвердить/снять, эскалировать owner | до 15 минут |
| Ошибка теста восстановления | Перезапустить тест, проверить практические действия и SLA | до 30 минут |
| Ошибка mobile smoke | Сохранить результат, проверить mobile readiness и content ingestion | до 30 минут |
| UI не реагирует | Ctrl+F5, повторный вход, проверить runbook статусы | до 10 минут |

## 3) Runbook сценарии
### Утренний чек
1. Нажать `Утренний чек`.
2. Проверить `runbookActionStatus`, `restoreActionStatus`, `alertsActionStatus`, `mobileActionStatus`.
3. При отклонениях выгрузить JSON/CSV.

### Инцидент-проверка
1. Нажать `Инцидент-проверка`.
2. Проверить практические действия, алерты, backup history, mobile/content readiness.
3. Зафиксировать выгрузку и передать ответственным.

## 4) Чеклист передачи смены
- [ ] Все high-алерты обработаны.
- [ ] Статус теста восстановления проверен.
- [ ] Mobile/content readiness проверены.
- [ ] При отклонениях выгружен отчет (JSON/CSV).
- [ ] Переданы комментарии следующему администратору.
- [ ] Нажата кнопка `Выйти`.

## 5) Горячие клавиши
- `Alt+1..5` — переключение вкладок.
- `/` — фокус на поиске.
- `Esc` — закрытие модального окна.
"""


def build_docx(now: str, text: str, images: list[Path], out_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Инструкция администратора Алхимик v2", level=0)
    doc.add_paragraph(f"Актуально на: {now}")

    doc.add_heading("Скрин 1: Вход в систему", level=1)
    doc.add_picture(str(images[0]), width=Inches(6.2))
    doc.add_paragraph("Кнопка входа находится в правом верхнем углу. После входа открывается рабочая панель.")

    doc.add_heading("Скрин 2: Рабочая панель пользователей", level=1)
    doc.add_picture(str(images[1]), width=Inches(6.2))

    doc.add_heading("Скрин 3: Вкладка безопасности", level=1)
    doc.add_picture(str(images[2]), width=Inches(6.2))
    doc.add_paragraph("Здесь доступны Утренний чек, Инцидент-проверка, автообновление, экспорт и управление алертами.")

    doc.add_heading("Скрин 4: Инструкция в веб-админке", level=1)
    doc.add_picture(str(images[3]), width=Inches(6.2))

    doc.add_heading("SLA и действия", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Ситуация"
    hdr[1].text = "Действие"
    hdr[2].text = "SLA"
    rows = [
        ("High алерт", "Подтвердить/снять, комментарий, эскалация", "до 15 минут"),
        ("Ошибка backup test", "Перезапуск + проверка действий", "до 30 минут"),
        ("Ошибка mobile smoke", "Сохранить, проверить readiness", "до 30 минут"),
        ("UI не отвечает", "Ctrl+F5, вход заново", "до 10 минут"),
    ]
    for a, b, c in rows:
        r = table.add_row().cells
        r[0].text = a
        r[1].text = b
        r[2].text = c

    doc.add_heading("Чеклист передачи смены", level=1)
    for line in [
        "Обработаны high-алерты",
        "Проверен статус теста восстановления",
        "Проверены mobile/content readiness",
        "Выгружен отчет при отклонениях",
        "Передан комментарий следующей смене",
        "Выполнен выход из системы",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_page_break()
    doc.add_heading("One-Page Printable: передача смены", level=1)
    for line in [
        "[ ] High-алерты обработаны",
        "[ ] Backup test проверен",
        "[ ] Mobile/content readiness проверены",
        "[ ] Отчет JSON/CSV выгружен при отклонениях",
        "[ ] Комментарий следующей смене передан",
        "[ ] Выход из системы выполнен",
    ]:
        doc.add_paragraph(line)

    doc.save(str(out_path))


def build_one_page_html(now: str, out_path: Path) -> None:
    html = f"""<!doctype html>
<html lang="ru">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>Алхимик One-Page Checklist</title>
  <style>
    body {{ font-family: Arial, sans-serif; color: #1f2f35; margin: 20px; }}
    h1 {{ margin: 0 0 8px; }}
    .muted {{ color: #5e727a; margin-bottom: 12px; }}
    table {{ width: 100%; border-collapse: collapse; margin-bottom: 12px; }}
    th, td {{ border: 1px solid #c5d5dc; padding: 7px; text-align: left; font-size: 13px; }}
    .check {{ margin: 6px 0; font-size: 14px; }}
    @media print {{
      body {{ margin: 10mm; }}
      .noprint {{ display: none; }}
    }}
  </style>
</head>
<body>
  <h1>Алхимик: One-Page инструкция дежурного администратора</h1>
  <div class="muted">Актуально на: {now}</div>
  <table>
    <thead><tr><th>Ситуация</th><th>Действие</th><th>SLA</th></tr></thead>
    <tbody>
      <tr><td>High алерт</td><td>Подтвердить/снять + комментарий + эскалация</td><td>до 15 минут</td></tr>
      <tr><td>Ошибка backup test</td><td>Повторный запуск + проверка практических действий</td><td>до 30 минут</td></tr>
      <tr><td>Ошибка mobile smoke</td><td>Сохранить результат + проверить readiness</td><td>до 30 минут</td></tr>
      <tr><td>UI не реагирует</td><td>Ctrl+F5 + повторный вход</td><td>до 10 минут</td></tr>
    </tbody>
  </table>
  <h3>Чеклист передачи смены</h3>
  <div class="check">[ ] Обработаны high-алерты</div>
  <div class="check">[ ] Проверен тест восстановления</div>
  <div class="check">[ ] Проверены mobile/content readiness</div>
  <div class="check">[ ] При отклонениях выгружен отчет JSON/CSV</div>
  <div class="check">[ ] Комментарий передан следующему администратору</div>
  <div class="check">[ ] Выполнен выход из системы</div>
  <button class="noprint" onclick="window.print()">Печать</button>
</body>
</html>
"""
    out_path.write_text(html, encoding="utf-8")


def main() -> None:
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    images = capture_screenshots()
    text = build_text(now)

    md = TOOLS / "admin-manual-ru-v2.md"
    txt = TOOLS / "admin-manual-ru-v2.txt"
    docx = TOOLS / "admin-manual-ru-v2.docx"
    one_page_html = TOOLS / "admin-one-page-ru.html"

    md.write_text(text, encoding="utf-8")
    txt.write_text(text, encoding="utf-8")
    build_docx(now, text, images, docx)
    build_one_page_html(now, one_page_html)

    # Копии в web_admin, чтобы инструкции открывались прямо из интерфейса
    for src_name in [
        "admin-manual-ru-v2.md",
        "admin-manual-ru-v2.txt",
        "admin-manual-ru-v2.docx",
        "admin-one-page-ru.html",
    ]:
        src = TOOLS / src_name
        dst = WEB_ADMIN / src_name
        dst.write_bytes(src.read_bytes())

    print("ok", md)
    print("ok", txt)
    print("ok", docx)
    print("ok", one_page_html)
    print("screens", [p.name for p in images])


if __name__ == "__main__":
    main()
