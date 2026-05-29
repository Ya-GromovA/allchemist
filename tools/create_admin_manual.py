from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, ArrowStyle, FancyArrowPatch


base = Path("/root/synapse/tools")
assets = base / "admin-manual-assets"
assets.mkdir(parents=True, exist_ok=True)


def flow_image(path: Path, title: str, boxes: list[tuple[float, float, str]], arrows: list[tuple[tuple[float, float], tuple[float, float]]]) -> None:
    fig, ax = plt.subplots(figsize=(12, 6.75), dpi=160)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 60)
    ax.axis("off")
    fig.patch.set_facecolor("#f5f9fb")
    ax.text(2, 56, title, fontsize=16, fontweight="bold", color="#103744")
    for x, y, text in boxes:
        patch = FancyBboxPatch(
            (x, y),
            24,
            10,
            boxstyle="round,pad=0.5,rounding_size=2",
            facecolor="#ffffff",
            edgecolor="#76a9bb",
            linewidth=1.5,
        )
        ax.add_patch(patch)
        ax.text(x + 12, y + 5, text, ha="center", va="center", fontsize=10, color="#16313b", wrap=True)
    for (x1, y1), (x2, y2) in arrows:
        arr = FancyArrowPatch(
            (x1, y1),
            (x2, y2),
            arrowstyle=ArrowStyle("Simple", head_length=8, head_width=8, tail_width=1.0),
            color="#2f7e99",
            linewidth=1.0,
            alpha=0.9,
        )
        ax.add_patch(arr)
    plt.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def panel_image(path: Path, title: str, lines: list[str]) -> None:
    fig, ax = plt.subplots(figsize=(12, 6.75), dpi=160)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 60)
    ax.axis("off")
    fig.patch.set_facecolor("#eef5f8")
    ax.text(3, 55, title, fontsize=16, fontweight="bold", color="#11343f")
    card = FancyBboxPatch((3, 6), 94, 44, boxstyle="round,pad=1.0,rounding_size=2", facecolor="#ffffff", edgecolor="#9fc2cf", linewidth=1.5)
    ax.add_patch(card)
    y = 45
    for line in lines:
        ax.text(7, y, line, fontsize=11, color="#173842", va="top")
        y -= 6
    plt.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


img_login = assets / "01-login-flow.png"
img_daily = assets / "02-daily-check.png"
img_incident = assets / "03-incident-runbook.png"
img_user = assets / "04-user-lifecycle.png"

flow_image(
    img_login,
    "Схема 1. Вход администратора в интерфейс",
    [
        (4, 32, "Открыть /api/v1/admin/web"),
        (30, 32, "Кнопка \"Вход в систему\"\n(правый верхний угол)"),
        (56, 32, "Логин/пароль\nadmin/admin123\nили OTP по телефону"),
        (82, 32, "Рабочая панель\nвкладок"),
    ],
    [((28, 37), (30, 37)), ((54, 37), (56, 37)), ((80, 37), (82, 37))],
)

panel_image(
    img_daily,
    "Схема 2. Ежедневный утренний чек (без инцидента)",
    [
        "1) Открыть вкладку \"Безопасность\" и нажать \"Утренний чек\".",
        "2) Проверить inline-статусы: runbookActionStatus, restoreActionStatus, alertsActionStatus.",
        "3) Если есть алерты: зафиксировать комментарий и подтвердить/снять подтверждение.",
        "4) Если тест восстановления просрочен: запустить \"Запустить тест восстановления\".",
        "5) Сохранить результат mobile smoke (кнопка \"Сохранить результат проверки\").",
        "6) Выгрузить отчет безопасности в JSON/CSV для архива.",
    ],
)

panel_image(
    img_incident,
    "Схема 3. Инцидент-проверка (при жалобах/сбоях)",
    [
        "1) Нажать \"Инцидент-проверка\" на вкладке \"Безопасность\".",
        "2) Проверить \"Практические действия\" и SLA-критичные пункты.",
        "3) Проверить \"Алерты безопасности\": high -> обработать в первую очередь.",
        "4) Проверить статус и историю теста восстановления.",
        "5) Обновить мобильную готовность и готовность загрузки контента.",
        "6) Зафиксировать экспортом CSV/JSON и передать отчет владельцу/техлиду.",
    ],
)

flow_image(
    img_user,
    "Схема 4. Сценарий создания и сопровождения пользователя",
    [
        (4, 32, "Вкладка \"Пользователи\"\nСоздать пользователя"),
        (30, 32, "Роль + план + модуль\n(или ручной ввод)"),
        (56, 32, "Проверить\ncreateUserResult\nи таблицу"),
        (82, 32, "Вкладка \"Подписки\"\nвыдать/отозвать\nи проверить KPI"),
    ],
    [((28, 37), (30, 37)), ((54, 37), (56, 37)), ((80, 37), (82, 37))],
)

now = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
md_path = base / "admin-manual-ru.md"
txt_path = base / "admin-manual-ru.txt"

manual_text = f"""# Инструкция администратора Алхимик (RU)

Актуально на: {now}

## 1. Назначение
Документ описывает повседневную работу администратора в веб-кабинете Алхимик: вход, управление пользователями, подписками, правами, журналом и блоком безопасности.

## 2. Доступ в систему
- URL: `/api/v1/admin/web`
- Вход через кнопку `Вход в систему` в правом верхнем углу.
- Способы входа:
  - Логин/пароль администратора.
  - OTP по телефону (код подтверждения).
- Дефолтные учетные данные (если не переопределены ENV):
  - Логин: `admin`
  - Пароль: `admin123`

## 3. Базовый рабочий цикл (каждый день)
1. Выполнить вход.
2. Перейти во вкладку `Безопасность`.
3. Нажать `Утренний чек`.
4. Проверить статусы:
   - `runbookActionStatus`
   - `restoreActionStatus`
   - `alertsActionStatus`
   - `mobileActionStatus`
5. При необходимости выгрузить отчет (`JSON/CSV`) и заархивировать.

## 4. Что делать в разных ситуациях

### Ситуация A: Пустой экран / кнопка не работает
- Нажать `Ctrl+F5` (полная перезагрузка).
- Повторить вход через кнопку в шапке.
- Проверить, что отображаются вкладки `Пользователи`, `Подписки`, `Права доступа`, `Журнал`, `Безопасность`.

### Ситуация B: Нужно создать пользователя
- Вкладка `Пользователи` -> заполнить телефон, роль, при необходимости план/модуль.
- Нажать `Создать пользователя и проверить`.
- Убедиться, что `createUserResult` содержит `userId` и выбранную роль.
- При необходимости назначить роль кнопкой `Назначить роль выбранному`.

### Ситуация C: Нужно выдать/отозвать доступ
- Выбрать пользователя.
- Перейти в `Подписки`.
- Указать план/модуль и нажать `Выдать` или `Отозвать`.
- Проверить `subscriptionsActionStatus` и обновить KPI.

### Ситуация D: Есть алерты безопасности
- Вкладка `Безопасность` -> `Обновить алерты`.
- Выбрать код алерта и заполнить комментарий.
- Нажать `Подтвердить алерт` или `Снять подтверждение`.
- Проверить, что статус обновился и запись видна в таблице.

### Ситуация E: Инцидент / жалоба на нестабильность
- Нажать `Инцидент-проверка`.
- Проверить:
  - `Практические действия`
  - `Алерты безопасности`
  - `Состояние/историю теста восстановления`
  - `Готовность мобильного контура`
  - `Готовность загрузки контента`
- Сформировать выгрузку JSON/CSV и передать ответственным.

### Ситуация F: Нужно быстро сбросить фильтры
- Во вкладке `Безопасность` нажать `Сбросить фильтры`.
- Проверить, что период и фильтры критичности/подтверждения вернулись к значениям по умолчанию.

## 5. Горячие клавиши оператора
- `Alt+1..5` — быстрый переход между вкладками.
- `/` — фокус на поиске текущей вкладки.
- `Esc` — закрытие модального окна.

## 6. Рекомендации по эскалации
- `high` алерты: реакция немедленно, уведомить owner/tech lead.
- Ошибка теста восстановления: приоритет высокий, запуск повторного теста и проверка практических действий.
- Повторные сбои мобильного smoke: вынести в инцидент, зафиксировать заметку и время проверки.

## 7. Чеклист завершения смены
- Обработаны все `high` алерты.
- Подтверждены/сняты подтверждения с комментарием.
- Проверены backup/mobile/content summary.
- При необходимости выгружен отчет и передан ответственным.
- Выполнен выход (`Выйти` в правом верхнем углу).
"""

md_path.write_text(manual_text, encoding="utf-8")
txt_path.write_text(manual_text, encoding="utf-8")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

doc.add_heading("Инструкция администратора Алхимик", level=0)
doc.add_paragraph(f"Актуально на: {now}")

doc.add_heading("1. Доступ и вход", level=1)
doc.add_paragraph(
    "Откройте /api/v1/admin/web и нажмите кнопку \"Вход в систему\" в правом верхнем углу. Используйте вход по логину/паролю или по OTP."
)
doc.add_picture(str(img_login), width=Inches(6.7))

doc.add_heading("2. Ежедневный регламент", level=1)
doc.add_paragraph("Рекомендуется начинать смену с кнопки \"Утренний чек\" на вкладке Безопасность.")
doc.add_picture(str(img_daily), width=Inches(6.7))

doc.add_heading("3. Действия при инциденте", level=1)
doc.add_paragraph("При жалобах пользователей, ошибках мобильного контура или нестабильности запускайте \"Инцидент-проверка\".")
doc.add_picture(str(img_incident), width=Inches(6.7))

doc.add_heading("4. Пользователи и подписки", level=1)
doc.add_paragraph("Создайте пользователя, назначьте роль/план/модуль, затем проверьте выдачу доступа во вкладке Подписки.")
doc.add_picture(str(img_user), width=Inches(6.7))

doc.add_heading("5. Краткий справочник ситуаций", level=1)
for item in [
    "Кнопка не реагирует: выполнить Ctrl+F5, затем повторно нажать \"Вход в систему\".",
    "Есть high-алерт: обработать в первую очередь, оставить комментарий, уведомить owner/tech lead.",
    "Провален тест восстановления: запустить повторно, проверить практические действия и SLA.",
    "Нужно собрать отчет: выбрать формат JSON/CSV и выполнить выгрузку.",
    "Нужно быстро вернуть стандартный вид: нажать \"Сбросить фильтры\".",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("6. Горячие клавиши", level=1)
doc.add_paragraph("Alt+1..5 — вкладки, / — поиск, Esc — закрыть модалку.")

docx_path = base / "admin-manual-ru.docx"
doc.save(str(docx_path))

print("created", md_path)
print("created", txt_path)
print("created", docx_path)
print("images", sorted([p.name for p in assets.glob("*.png")]))
